import asyncio
from crawl4ai import AsyncWebCrawler
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from bs4 import BeautifulSoup
import os
import re

class Crawl4AiDocScraper:
    def __init__(self):
        self.doc = Document()
        self.setup_document()
        self.urls = {
            "Installation": "https://crawl4ai.com/mkdocs/basic/installation/",
            "Docker Deployment": "https://crawl4ai.com/mkdocs/basic/docker-deploymeny/",
            "Quick Start": "https://crawl4ai.com/mkdocs/basic/quickstart/",
            "Simple Crawling": "https://crawl4ai.com/mkdocs/basic/simple-crawling/",
            "Output Formats": "https://crawl4ai.com/mkdocs/basic/output-formats/",
            "Browser Configuration": "https://crawl4ai.com/mkdocs/basic/browser-config/",
            "Page Interaction": "https://crawl4ai.com/mkdocs/basic/page-interaction/",
            "Content Selection": "https://crawl4ai.com/mkdocs/basic/content-selection/",
            "Content Processing": "https://crawl4ai.com/mkdocs/advanced/content-processing/",
            "Magic Mode": "https://crawl4ai.com/mkdocs/advanced/magic-mode/",
            "Hooks & Auth": "https://crawl4ai.com/mkdocs/advanced/hooks-auth/",
            "Proxy & Security": "https://crawl4ai.com/mkdocs/advanced/proxy-security/",
            "Session Management": "https://crawl4ai.com/mkdocs/advanced/session-management/",
            "Advanced Session Management": "https://crawl4ai.com/mkdocs/advanced/session-management-advanced/",
            "Extraction Overview" : "https://crawl4ai.com/mkdocs/extraction/overview/",
            "LLM Strategy" : "https://crawl4ai.com/mkdocs/extraction/llm/",
            "Json-CSS Extractor Basic" : "https://crawl4ai.com/mkdocs/extraction/css/",
            "Json-CSS Extractor Advanced" : "https://crawl4ai.com/mkdocs/extraction/css-advanced/",
            "Cosine Strategy" : "https://crawl4ai.com/mkdocs/extraction/cosine/",
            "Chunking" : "https://crawl4ai.com/mkdocs/extraction/chunking/",
            "Parameters Table" : "https://crawl4ai.com/mkdocs/api/parameters/",
            "AsyncWebCrawler" : "https://crawl4ai.com/mkdocs/api/async-webcrawler/",
            "AsyncWebCrawler.arun()" : "https://crawl4ai.com/mkdocs/api/arun/",
            "CrawlResult" : "CrawlResult",
            "Strategies" : "https://crawl4ai.com/mkdocs/api/strategies/",

        }

    def setup_document(self):
        # Document styles setup
        title = self.doc.add_heading('Crawl4AI Documentation', level=0)
        font = title.runs[0].font
        font.size = Pt(20)
        font.color.rgb = RGBColor(31, 73, 125)
        
        # Code style
        style = self.doc.styles.add_style('Code', 1)
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(10)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        
        # Note style
        note_style = self.doc.styles.add_style('Note', 1)
        font = note_style.font
        font.italic = True
        font.size = Pt(10)

    def clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        text = re.sub(r'\s+', ' ', text)
        text = text.replace('\n', ' ').strip()
        return text

    def process_html_content(self, html: str):
        """Process HTML content with improved debugging"""
        try:
            soup = BeautifulSoup(html, 'html.parser')
            print(f"HTML length: {len(html)}")
            
            # Try different possible content containers
            content_containers = [
                soup.find('div', class_='md-content'),
                soup.find('article'),
                soup.find('main'),
                soup.find('div', class_='content'),
                soup.find('body')  # Fallback
            ]
            
            main_content = next((container for container in content_containers if container), None)
            
            if not main_content:
                print("No content container found")
                return None
                
            print(f"Found content container: {main_content.name}")
            
            content = []
            # Get all relevant elements
            elements = main_content.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'pre', 'ul', 'ol', 'code'])
            
            for element in elements:
                # Process headings
                if element.name.startswith('h') and len(element.name) == 2:
                    text = self.clean_text(element.get_text())
                    if text:
                        content.append({
                            'type': 'heading',
                            'level': int(element.name[1]),
                            'text': text
                        })
                
                # Process code blocks
                elif element.name == 'pre' or element.name == 'code':
                    code_text = element.get_text()
                    if code_text.strip():
                        # Try to detect language from class
                        classes = element.get('class', [])
                        language = next((cls.replace('language-', '') 
                                      for cls in classes 
                                      if cls.startswith('language-')), '')
                        content.append({
                            'type': 'code',
                            'language': language,
                            'text': code_text.strip()
                        })
                
                # Process lists
                elif element.name in ['ul', 'ol']:
                    items = [self.clean_text(li.get_text()) 
                            for li in element.find_all('li')]
                    if items:
                        content.append({
                            'type': 'list',
                            'items': items
                        })
                
                # Process paragraphs
                elif element.name == 'p':
                    text = self.clean_text(element.get_text())
                    if text:
                        content.append({
                            'type': 'text',
                            'text': text
                        })
            
            print(f"Extracted {len(content)} content blocks")
            return content
            
        except Exception as e:
            print(f"Error processing HTML: {str(e)}")
            return None

    async def extract_section_content(self, title: str, url: str):
        """Extract content with improved error handling"""
        async with AsyncWebCrawler() as crawler:
            try:
                print(f"\nFetching {url}...")
                result = await crawler.arun(
                    url=url,
                    magic=True,
                    delay_before_return_html=3.0,  # Increased delay
                    process_iframes=True,
                    screenshot=True  # For debugging
                )
                
                if not result.success:
                    print(f"Crawl failed: {result.error_message}")
                    return None
                
                print("Successfully fetched page")
                content = self.process_html_content(result.html)
                
                if not content:
                    print("No content extracted")
                    return None
                    
                return {
                    'title': title,
                    'content': content
                }
                
            except Exception as e:
                print(f"Error during extraction: {str(e)}")
                return None

    def add_content_to_doc(self, content):
        """Add content to document with improved formatting"""
        if not content:
            return
            
        # Add section title
        title = self.doc.add_heading(content['title'], level=1)
        title.runs[0].font.color.rgb = RGBColor(31, 73, 125)
        
        for item in content['content']:
            if item['type'] == 'heading':
                heading = self.doc.add_heading(item['text'], level=item['level'])
                heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)
                
            elif item['type'] == 'code':
                p = self.doc.add_paragraph(style='Code')
                if item['language']:
                    p.add_run(f"Language: {item['language']}\n").bold = True
                p.add_run(item['text'])
                
            elif item['type'] == 'list':
                for idx, list_item in enumerate(item['items'], 1):
                    self.doc.add_paragraph(f"{idx}. {list_item}")
                    
            elif item['type'] == 'text':
                self.doc.add_paragraph(item['text'])

    async def scrape_all(self):
        """Scrape all documentation sections"""
        for title, url in self.urls.items():
            print(f"\nScraping {title}...")
            content = await self.extract_section_content(title, url)
            if content:
                self.add_content_to_doc(content)
                print(f"✓ Successfully processed {title}")
            else:
                print(f"✗ Failed to process {title}")
            await asyncio.sleep(3)  # Increased delay between requests

    def save_doc(self, filename="Crawl4AI_Documentation.docx"):
        """Save the document"""
        os.makedirs('docs', exist_ok=True)
        path = os.path.join('docs', filename)
        self.doc.save(path)
        print(f"\nDocumentation saved to {path}")

async def main():
    scraper = Crawl4AiDocScraper()
    print("Starting documentation scraping...")
    await scraper.scrape_all()
    scraper.save_doc()
    print("Scraping completed!")

if __name__ == "__main__":
    import time
    start_time = time.time()
    asyncio.run(main())
    print(f"\nTotal execution time: {time.time() - start_time:.2f} seconds")